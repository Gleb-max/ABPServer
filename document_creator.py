from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.shared import Inches, Cm, Pt

from data.dean import Dean
from data.enrollee import Enrollee
from data.study_direction import StudyDirection
from datetime import datetime

from data.user import User
from docx2pdf import convert as convert_docx2pdf


def create_order_of_admission(filename, enrolls: List[Enrollee], direction: StudyDirection, group_number):
    document = Document()

    upper = document.add_paragraph('Федеральное государственное бюджетное образовательное учреждение\n'
                                   'высшего образования\n'
                                   'Сызранский государственный университет имени Филиппа Лимонадова\n\n')
    upper.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = document.add_paragraph()
    header.add_run('Приказ').bold = True
    header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        'На основании решения приемной комиссии зачислить на 1 курс очной формы обучения с 01.09.2021 г.\n\n')

    paragraph = document.add_paragraph(f'Факультет “{direction.faculty.name}”\n '
                                       f'Направление “{direction.name}”\n'
                                       f'Номер группы присвоить “{group_number}”\n')
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=len(enrolls) + 1, cols=4, style='Table Grid')
    table.cell(0, 0).text = '№'
    table.cell(0, 1).text = 'ФИО'
    table.cell(0, 2).text = 'Основа'
    table.cell(0, 3).text = 'Баллы'

    for row_num, enroll in enumerate(enrolls, 1):
        print(f'{enroll.user.name} was written to table')
        table.cell(row_num, 0).text = str(row_num)
        table.cell(row_num, 1).text = f'{enroll.user.name} {enroll.user.surname} {enroll.user.last_name}'
        table.cell(row_num, 2).text = 'бюджет' if enroll.is_budgetary else 'платная'
        table.cell(row_num, 3).text = str(enroll.get_total_grade)

    document.add_paragraph('\n\n\nРектор\t\t\t\t\t\t\t\t\t Иванов М.Ю.')

    full_file_name = f'{filename}.docx'.replace(' ', '_')
    document.save(full_file_name)
    return full_file_name


def get_tabs(_len, data, field_name):
    data = '  ' + (data if data else ' ')
    return data + int((_len - len(field_name) - len(data)) // 10) * '\t'


def format_datetime(date: datetime):
    if date:
        return date.strftime('%d.%M.%Y')
    return ' '


def create_student_personal_profile(filename, user: User, need_pdf=False):
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # crutch to rotate
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    upper = document.add_paragraph('Федеральное государственное бюджетное образовательное учреждение '
                                   'высшего образования\n'
                                   'Сызранский государственный университет имени Филиппа Лимонадова\n')
    upper.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    personal_debt = user.student.card_number

    delo = document.add_paragraph(f'Личное дело №')
    delo.add_run(f'{personal_debt}').underline = True
    delo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    line_len = 108

    par = document.add_paragraph()
    study_direction = user.student.student_group.direction

    par.add_run('Факультет')
    facult = study_direction.faculty.name
    par.add_run(get_tabs(line_len, facult, 'Факультет') + '\n').underline = True
    par.add_run('Специальность')
    spec = study_direction.name
    par.add_run(get_tabs(line_len, spec, 'Специальность') + '\n').underline = True
    par.add_run('Группа')
    group = user.student.student_group.name
    par.add_run(get_tabs(line_len // 1.5, group, 'Группа')).underline = True

    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    header = document.add_paragraph('Учебная карточка студента\n(очная форма)')
    header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tab = document.add_table(rows=1, cols=2)
    tab.autofit = False
    tab.allow_autofit = False
    # par = document.add_paragraph()
    cur_cell = tab.cell(0, 0)
    cur_cell.width = Inches(6.5)
    par = cur_cell.add_paragraph()
    par.add_run('Основа')  # TODO filing
    small_line = '\t\t\t\t\t'
    par.add_run(small_line).underline = True

    par.add_run(', личное дело №')
    par.add_run(f'{personal_debt}\n').underline = True

    par.add_run('ФИО')
    fio = f'  {user.name} {user.surname} {user.last_name}'
    par.add_run(get_tabs(line_len, fio, 'ФИО') + '\n').underline = True

    par.add_run('Дата рождения')
    par.add_run(get_tabs(line_len, format_datetime(user.enrollee.birthday), 'Дата рождения') + '\n').underline = True

    par.add_run('Место рождения')
    par.add_run(get_tabs(line_len, user.enrollee.birth_place, 'Место рождения') + '\n').underline = True
    par.add_run('Паспорт')
    passport = user.enrollee.passport
    passport_info = f' {passport.series} {passport.number}'
    par.add_run(passport_info).underline = True
    par.add_run(' выдан')

    vidan = passport.who_issued + str(passport.department_code)
    par.add_run(get_tabs(line_len, vidan, passport_info + 'паспорт  выдан') + '\n').underline = True

    small_line_len = 70
    par.add_run('\n\nКонтакты\n').bold = True
    par.add_run('Телефон')
    par.add_run(get_tabs(line_len // 1.5, user.enrollee.phone, 'Телефон') + '\n').underline = True
    par.add_run('E-mail')
    par.add_run(get_tabs(line_len // 1.5, user.email, 'e-mail') + '\n').underline = True
    par.add_run('Адрес по прописке')
    par.add_run(get_tabs(line_len, passport.registration_address, 'e-mail') + '\n').underline = True
    par.add_run(get_tabs(line_len, '', '') + '\t' + '\n').underline = True
    par.add_run(get_tabs(line_len, '', '') + '\t' + '\n').underline = True

    import os.path

    if user.enrollee.photo and os.path.isfile(user.enrollee.photo):
        img_path = user.enrollee.photo
    else:
        img_path = 'default_photo.jpg'

    with open(img_path, 'rb') as img:
        cur_cell = tab.cell(0, 1)
        paragraph = cur_cell.paragraphs[0]
        run = paragraph.add_run()
        picture = run.add_picture(img, width=Inches(1.5), height=Inches(1.5))
        cur_cell.width = Cm(10)

    input_file_path = f'{filename}.docx'
    document.save(input_file_path)

    if need_pdf:
        out_file_path = f'{filename}.pdf'
        import pythoncom
        pythoncom.CoInitializeEx(0)
        convert_docx2pdf(input_file_path, out_file_path)
        return out_file_path

    return input_file_path


def create_student_record_book(filename, users: List[User], need_pdf=False):
    document = Document()

    for ind, user in enumerate(users):
        # section = document.add_section(WD_SECTION.NEW_PAGE)
        section = document.sections[-1]
        section.top_margin = Cm(0.7)
        section.orientation = WD_ORIENT.LANDSCAPE
        # crutch to rotate
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        upper = document.add_paragraph('Зачетная книжка №')
        upper.add_run(f'{user.student.record_book_number}').underline = True
        upper.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        line_len = 60

        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        # par = document.add_paragraph()
        cur_cell = table.cell(0, 1)
        cur_cell.width = Inches(6.5)
        par = cur_cell.add_paragraph()
        par.add_run('ФИО')
        fio = f'  {user.name} {user.surname} {user.last_name}'
        par.add_run(get_tabs(line_len, fio, 'ФИО') + '\n').underline = True
        par.add_run('Факультет')
        directional = user.enrollee.study_direction
        par.add_run(get_tabs(line_len, directional.faculty.name, 'Факультет') + '\n').underline = True
        par.add_run('Специальность')
        par.add_run(get_tabs(line_len, directional.name, 'Специальность') + '\n').underline = True
        par.add_run(get_tabs(line_len, '', '') + '\n\n').underline = True
        par.add_run(f'Поступил в {user.student.enrollment_date.year} году')

        import os.path
        if user.enrollee.photo and os.path.isfile(user.enrollee.photo):
            img_path = user.enrollee.photo
        else:
            img_path = 'default_photo.jpg'

        with open(img_path, 'rb') as img:
            cur_cell = table.cell(0, 0)
            cur_cell.width = Cm(6.2)
            paragraph = cur_cell.paragraphs[0]
            run = paragraph.add_run()
            picture = run.add_picture(img, width=Inches(1.5), height=Inches(1.5))

        workers = Dean.query.filter_by(faculty_pk=directional.id)

        prorector = workers.filter(Dean.post == 'Проректор').first()
        dean = workers.filter(Dean.post == 'Декан').first()

        paragraph = document.add_paragraph()
        fio = ' ...'
        if prorector:
            fio = f'{prorector.user.surname} {prorector.user.name[:1]}.{prorector.user.last_name[:1]}.'

        paragraph.add_run(f'\n\nПроректор по учебной работе {fio}\n\n')

        fio = ' ...'
        if prorector:
            fio = f'{dean.user.surname} {dean.user.name[:1]}.{dean.user.last_name[:1]}.'
        paragraph.add_run(f'Проректор по учебной работе {fio}\n')

        if ind % 2 == 0:
            paragraph.add_run('-' * 160 + '\n')

    input_file_path = f'{filename}.docx'
    document.save(input_file_path)

    if need_pdf:
        out_file_path = f'{filename}.pdf'
        import pythoncom
        pythoncom.CoInitializeEx(0)
        convert_docx2pdf(input_file_path, out_file_path)
        return out_file_path

    return input_file_path


def create_student_card(filename, users: List[User], need_pdf=False):
    document = Document()

    for (ind, user) in enumerate(users):
        section = document.sections[-1]
        section.top_margin = Cm(0.7)
        section.orientation = WD_ORIENT.LANDSCAPE
        # crutch to rotate
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        table = document.add_table(rows=1, cols=2, style='Table Grid')
        table.autofit = False
        table.allow_autofit = False

        # 1 col
        current_cell = table.cell(0, 0)
        current_cell.width = Cm(11)
        upper = current_cell.add_paragraph('Студенческий билет №')
        upper.add_run(f'{user.student.card_number}').underline = True
        upper.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        min_tabble = current_cell.add_table(rows=1, cols=2)
        # min_tabble.autofit = True
        # min_tabble.allow_autofit = True

        # photo
        min_tabble.cell(0, 0).width = Cm(4)
        paragraph = min_tabble.cell(0, 0).paragraphs[0]
        run = paragraph.add_run('\n')
        import os.path
        if user.enrollee.photo and os.path.isfile(user.enrollee.photo):
            img_path = user.enrollee.photo
        else:
            img_path = 'default_photo.jpg'

        with open(img_path, 'rb') as img:
            picture = run.add_picture(img, width=Inches(1.5), height=Inches(1.5))

        paragraph.add_run('\n\nДекан факультета')

        min_tabble.cell(0, 1).width = Cm(5.5)
        paragraph = min_tabble.cell(0, 1).add_paragraph()
        line_len = 35
        paragraph.add_run('Фамилия')
        paragraph.add_run(get_tabs(line_len, user.surname, 'Фамилия') + '\n').underline = True
        paragraph.add_run('Имя')
        paragraph.add_run(get_tabs(line_len, user.name, 'Имя') + '\n').underline = True
        paragraph.add_run('Отчество')
        paragraph.add_run(get_tabs(line_len, user.last_name, 'Отчество  ') + '\n').underline = True
        faculty = user.enrollee.study_direction.faculty.name
        paragraph.add_run('Факультет')
        paragraph.add_run(get_tabs(line_len, faculty, 'Факультет') + '\n').underline = True
        paragraph.add_run('Группа')
        paragraph.add_run(get_tabs(line_len, user.student.student_group.name, 'Группа') + '\n\n').underline = True
        paragraph.add_run('Проректор по учебной работе' + '\n\n')

        # 2 col
        current_cell = table.cell(0, 1)
        start_year = user.student.enrollment_date.year
        digids = start_year % 100
        paragraph = current_cell.add_paragraph('\n')

        font_size = 15
        for i in range(5):
            run = paragraph.add_run(f'В 20')
            run.font.size = Pt(font_size)
            run = paragraph.add_run(f'{digids + i}')
            run.underline = True
            run.font.size = Pt(font_size)

            paragraph.add_run('/')
            run = paragraph.add_run(f'{digids + i + 1} ')
            run.underline = True
            run.font.size = Pt(font_size)
            run = paragraph.add_run('г. является студентом ')
            run.font.size = Pt(font_size)
            run = paragraph.add_run(f'{i + 1}')
            run.underline = True
            run.font.size = Pt(font_size)
            run = paragraph.add_run(' курса\n\n')
            run.font.size = Pt(font_size)

    input_file_path = f'{filename}.docx'
    document.save(input_file_path)
    import os
    print('path', os.path.abspath(os.getcwd()))

    if need_pdf:
        out_file_path = f'{filename}.pdf'
        import pythoncom
        pythoncom.CoInitializeEx(0)
        convert_docx2pdf(input_file_path, out_file_path)
        return out_file_path

    return input_file_path


def create_instruct_table(filename, users: List[User], subject_name, need_pdf=False):
    document = Document()

    header = document.add_paragraph()
    header.add_run(f'Инструктаж по технике безопасности по предмету').bold = True
    header.add_run(f'  {subject_name} ').underline = True
    header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(cols=4, rows=len(users) + 2, style='Table Grid')

    table.rows[0].cells[0].merge(table.cell(1, 0))
    table.cell(0, 0).add_paragraph('№\nп/п')
    table.cell(1, 0).width = Cm(1.2)
    table.cell(2, 0).width = Cm(1.2)
    table.rows[0].cells[1].merge(table.cell(1, 1))
    p = table.cell(0, 1).add_paragraph('ФИО инструктируемого')
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.rows[0].cells[2].merge(table.cell(0, 3))
    table.cell(0, 2).add_paragraph('Группа')
    table.cell(1, 2).add_paragraph('Дата инструктажа')
    table.cell(1, 3).add_paragraph('Подпись')

    table.autofit = False
    table.allow_autofit = False
    norm_date = lambda x: x.strftime('%d.%m.%Y')

    for ind, user in enumerate(users, 2):
        numb_cell = table.cell(ind, 0)
        numb_cell.width = Cm(1.2)
        p = numb_cell.add_paragraph(str(ind - 1))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        numb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        fio_cell = table.cell(ind, 1)
        fio_cell.width = Inches(2.6)
        p = fio_cell.add_paragraph(f'{user.name} {user.surname} {user.last_name}')
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fio_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        date_cell = table.cell(ind, 2)
        date_cell.width = Inches(1.3)

        if user.student:
            p = date_cell.add_paragraph(f'{norm_date(user.student.enrollment_date)}')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        podp_cell = table.cell(ind, 3)
        podp_cell.width = Inches(1.3)

    par = document.add_paragraph("\n\nИнструктаж провел")
    par.add_run('\t\t\t\t\t\t').underline = True
    par.add_run('\t').underline = False
    par.add_run('\t\t\t').underline = True
    par = document.add_paragraph()
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = par.add_run('Дата, подпись')
    font = run.font
    font.size = Pt(8)

    input_file_path = f'{filename}.docx'
    document.save(input_file_path)

    if need_pdf:
        out_file_path = f'{filename}.pdf'
        import pythoncom
        pythoncom.CoInitializeEx(0)
        convert_docx2pdf(input_file_path, out_file_path)
        return out_file_path

    return input_file_path


def create_attendance_log():
    document = Document()

    header = document.add_paragraph()
    header.add_run(f'Журнал посещаемости группы').bold = True
    header.add_run(f'  {"subject_name"} ').underline = True
    header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(cols=17, rows=10 + 2, style='Table Grid')

    table.rows[0].cells[0].merge(table.cell(1, 0))
    table.cell(0, 0).add_paragraph('№\nп/п')
    table.cell(1, 0).width = Cm(1.2)
    table.cell(2, 0).width = Cm(1.2)

    table.rows[0].cells[1].merge(table.cell(1, 1))
    p = table.cell(0, 1).add_paragraph('ФИО обучающегося \ Месяц, число')
    table.cell(0, 1).width = Cm(5)
    table.cell(0, 2).width = Cm(5)
    p.width = Cm(5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    section = document.sections[-1]
    section.top_margin = Cm(0.7)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    table.rows[0].cells[2].merge(table.cell(0, 16))
    table.cell(0, 3).add_paragraph('Предмет')

    for i in range(2, 17):
        table.cell(1, i).add_paragraph(f'{i - 1}.09')
        tc = table.cell(1, i)._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), "btLr")  # btLr tbRl
        tcPr.append(textDirection)

        table.cell(1, i).height = Cm(2)
        table.cell(1, i).width = Cm(0.7)

    # table.rows[0].cells[2].merge(table.cell(0, 3))
    # table.cell(0, 2).add_paragraph('Группа')
    # table.cell(1, 2).add_paragraph('Дата инструктажа')
    # table.cell(1, 3).add_paragraph('Подпись')

    table.autofit = False
    table.allow_autofit = False
    norm_date = lambda x: x.strftime('%d.%m.%Y')

    # for ind, user in enumerate(users, 2):
    #     numb_cell = table.cell(ind, 0)
    #     numb_cell.width = Cm(1.2)
    #     p = numb_cell.add_paragraph(str(ind - 1))
    #     p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     numb_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #
    #     fio_cell = table.cell(ind, 1)
    #     fio_cell.width = Inches(2.6)
    #     p = fio_cell.add_paragraph(f'{user.name} {user.surname} {user.last_name}')
    #     p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     fio_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #
    #     date_cell = table.cell(ind, 2)
    #     date_cell.width = Inches(1.3)
    #
    #     if user.student:
    #         p = date_cell.add_paragraph(f'{norm_date(user.student.enrollment_date)}')
    #         p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #         date_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    #
    #     podp_cell = table.cell(ind, 3)
    #     podp_cell.width = Inches(1.3)

    par = document.add_paragraph("\n\nИнструктаж провел")
    par.add_run('\t\t\t\t\t\t').underline = True
    par.add_run('\t').underline = False
    par.add_run('\t\t\t').underline = True
    par = document.add_paragraph()
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = par.add_run('Дата, подпись')
    font = run.font
    font.size = Pt(8)

    input_file_path = f'{"123"}.docx'
    document.save(input_file_path)


    return input_file_path

create_attendance_log()